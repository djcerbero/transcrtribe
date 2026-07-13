"""Export conversation turns to txt, rtf, docx and pdf."""
from __future__ import annotations

from pathlib import Path

from .conversation import ConversationTurn, format_timestamp

SUPPORTED_FORMATS = ("txt", "rtf", "docx", "pdf")


def export(turns: list[ConversationTurn], output_base: Path, formats: list[str], title: str) -> list[Path]:
    written: list[Path] = []
    for fmt in formats:
        if fmt == "txt":
            written.append(_export_txt(turns, output_base.with_suffix(".txt")))
        elif fmt == "rtf":
            written.append(_export_rtf(turns, output_base.with_suffix(".rtf"), title))
        elif fmt == "docx":
            written.append(_export_docx(turns, output_base.with_suffix(".docx"), title))
        elif fmt == "pdf":
            written.append(_export_pdf(turns, output_base.with_suffix(".pdf"), title))
        else:
            raise ValueError(f"Unsupported export format: {fmt}")
    return written


def _export_txt(turns: list[ConversationTurn], path: Path) -> Path:
    lines = []
    for turn in turns:
        lines.append(f"[{format_timestamp(turn.start)}] {turn.speaker}: {turn.text}")
    path.write_text("\n\n".join(lines) + "\n", encoding="utf-8")
    return path


def _rtf_escape(text: str) -> str:
    text = text.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}")
    return text.encode("ascii", "xmlcharrefreplace").decode("ascii").replace("&#", r"\u").replace(";", " ")


def _export_rtf(turns: list[ConversationTurn], path: Path, title: str) -> Path:
    body = [r"{\rtf1\ansi\deff0", r"{\fonttbl{\f0 Helvetica;}}", r"\f0\fs24"]
    body.append(r"{\b " + _rtf_escape(title) + r"}\par\par")
    for turn in turns:
        header = f"[{format_timestamp(turn.start)}] {turn.speaker}:"
        body.append(r"{\b " + _rtf_escape(header) + r"} " + _rtf_escape(turn.text) + r"\par\par")
    body.append("}")
    path.write_text("\n".join(body), encoding="ascii", errors="ignore")
    return path


def _export_docx(turns: list[ConversationTurn], path: Path, title: str) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for turn in turns:
        p = doc.add_paragraph()
        run = p.add_run(f"[{format_timestamp(turn.start)}] {turn.speaker}: ")
        run.bold = True
        p.add_run(turn.text)
    doc.save(str(path))
    return path


def _export_pdf(turns: list[ConversationTurn], path: Path, title: str) -> Path:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    styles = getSampleStyleSheet()
    speaker_style = ParagraphStyle(
        "Speaker", parent=styles["Normal"], fontName="Helvetica-Bold", spaceAfter=2
    )
    text_style = ParagraphStyle("Body", parent=styles["Normal"], spaceAfter=12, leading=15)

    doc = SimpleDocTemplate(str(path), pagesize=LETTER, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.25 * inch)]
    for turn in turns:
        header = f"[{format_timestamp(turn.start)}] {turn.speaker}:"
        story.append(Paragraph(_escape_html(header), speaker_style))
        story.append(Paragraph(_escape_html(turn.text), text_style))
    doc.build(story)
    return path


def _escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
